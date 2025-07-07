import io
from datetime import datetime
from typing import Tuple
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import logging

logger = logging.getLogger(__name__)


def generate_docx(proposal) -> Tuple[bytes, str]:
    """
    Generate a DOCX file from proposal content
    
    Args:
        proposal: Proposal model object
        
    Returns:
        Tuple of (file_content_bytes, filename)
    """
    try:
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_heading(proposal.title, 0)
        title.alignment = 1  # Center alignment
        
        # Add metadata
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Proposal ID: {proposal.id}")
        doc.add_paragraph()
        
        # Add executive summary if available
        if proposal.executive_summary:
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(proposal.executive_summary)
            doc.add_paragraph()
        
        # Add main content
        doc.add_heading('Proposal Content', level=1)
        
        # Split content into paragraphs and add them
        content_paragraphs = proposal.content.split('\n\n')
        for paragraph in content_paragraphs:
            if paragraph.strip():
                # Check if it's a heading (starts with #)
                if paragraph.strip().startswith('#'):
                    heading_text = paragraph.strip().lstrip('# ')
                    doc.add_heading(heading_text, level=2)
                else:
                    doc.add_paragraph(paragraph.strip())
        
        # Add footer with proposal metadata
        doc.add_page_break()
        doc.add_heading('Proposal Information', level=1)
        doc.add_paragraph(f"Status: {proposal.status.title()}")
        doc.add_paragraph(f"Version: {proposal.version}")
        if proposal.confidence_score:
            doc.add_paragraph(f"Confidence Score: {proposal.confidence_score:.2f}")
        if proposal.user_rating:
            doc.add_paragraph(f"User Rating: {proposal.user_rating}/5 stars")
        
        # Save to BytesIO
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Generate filename
        safe_title = "".join(c for c in proposal.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        filename = f"{safe_title[:50]}_proposal.docx"
        
        return file_stream.getvalue(), filename
        
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        raise


def generate_pdf(proposal) -> Tuple[bytes, str]:
    """
    Generate a PDF file from proposal content
    
    Args:
        proposal: Proposal model object
        
    Returns:
        Tuple of (file_content_bytes, filename)
    """
    try:
        # Create PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        
        # Add title
        pdf.cell(0, 10, proposal.title, ln=True, align='C')
        pdf.ln(10)
        
        # Add metadata
        pdf.set_font('Arial', '', 10)
        pdf.cell(0, 5, f"Generated on: {datetime.now().strftime('%B %d, %Y')}", ln=True)
        pdf.cell(0, 5, f"Proposal ID: {str(proposal.id)}", ln=True)
        pdf.ln(10)
        
        # Add executive summary if available
        if proposal.executive_summary:
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 8, 'Executive Summary', ln=True)
            pdf.set_font('Arial', '', 11)
            
            # Split text to fit in PDF width
            summary_lines = proposal.executive_summary.split('\n')
            for line in summary_lines:
                # Handle long lines by wrapping them
                words = line.split(' ')
                current_line = ""
                for word in words:
                    test_line = current_line + word + " "
                    if pdf.get_string_width(test_line) < 180:  # Max width
                        current_line = test_line
                    else:
                        if current_line:
                            pdf.cell(0, 6, current_line.strip(), ln=True)
                        current_line = word + " "
                if current_line:
                    pdf.cell(0, 6, current_line.strip(), ln=True)
            pdf.ln(5)
        
        # Add main content
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 8, 'Proposal Content', ln=True)
        pdf.set_font('Arial', '', 11)
        
        # Process content paragraphs
        content_paragraphs = proposal.content.split('\n\n')
        for paragraph in content_paragraphs:
            if paragraph.strip():
                # Check if it's a heading
                if paragraph.strip().startswith('#'):
                    heading_text = paragraph.strip().lstrip('# ')
                    pdf.set_font('Arial', 'B', 12)
                    pdf.ln(3)
                    pdf.cell(0, 7, heading_text, ln=True)
                    pdf.set_font('Arial', '', 11)
                else:
                    # Handle paragraph wrapping
                    lines = paragraph.strip().split('\n')
                    for line in lines:
                        words = line.split(' ')
                        current_line = ""
                        for word in words:
                            test_line = current_line + word + " "
                            if pdf.get_string_width(test_line) < 180:
                                current_line = test_line
                            else:
                                if current_line:
                                    pdf.cell(0, 6, current_line.strip(), ln=True)
                                current_line = word + " "
                        if current_line:
                            pdf.cell(0, 6, current_line.strip(), ln=True)
                    pdf.ln(3)
        
        # Add footer with metadata
        pdf.add_page()
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 8, 'Proposal Information', ln=True)
        pdf.set_font('Arial', '', 11)
        pdf.cell(0, 6, f"Status: {proposal.status.title()}", ln=True)
        pdf.cell(0, 6, f"Version: {proposal.version}", ln=True)
        if proposal.confidence_score:
            pdf.cell(0, 6, f"Confidence Score: {proposal.confidence_score:.2f}", ln=True)
        if proposal.user_rating:
            pdf.cell(0, 6, f"User Rating: {proposal.user_rating}/5 stars", ln=True)
        
        # Generate filename
        safe_title = "".join(c for c in proposal.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        filename = f"{safe_title[:50]}_proposal.pdf"
        
        return pdf.output(dest='S').encode('latin-1'), filename
        
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        raise 